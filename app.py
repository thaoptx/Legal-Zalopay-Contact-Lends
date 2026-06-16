#!/usr/bin/env python3
"""
Zalopay Legal Contract — Rà soát Hợp đồng Cộng tác viên
GreenNode Claw-a-thon 2026
"""

import os, sys, json, tempfile, time
from pathlib import Path
from http.server import HTTPServer, BaseHTTPRequestHandler
from openai import OpenAI

PORT = 8080
SCRIPT_DIR = Path(__file__).parent

# ── GreenNode MaaS config ──────────────────────────────
GN_API_KEY  = os.environ.get("AI_PLATFORM_API_KEY", "vn-mn0D-2X-12U8_XoVAIqA_XxuUxmVy-52872d7d3566478f90391e9202b2d2f8_o6llCb8_W_G0rTKL5eyqH34SQOpd-y")
GN_BASE_URL = "https://maas-llm-aiplatform-hcm.api.vngcloud.vn/v1"
GN_MODEL    = "minimax/minimax-m2.5"      # Main review model
GN_MODEL_JSON = "qwen/qwen3-5-27b"        # For structured JSON responses
SYSTEM_PROMPT_FILE = SCRIPT_DIR / "system_prompt.txt"
DATA_FILE = SCRIPT_DIR / "projects.json"

def load_system_prompt():
    if SYSTEM_PROMPT_FILE.exists():
        return SYSTEM_PROMPT_FILE.read_text(encoding="utf-8")
    return "Bạn là chuyên gia rà soát hợp đồng giao khoán."

def load_projects():
    if DATA_FILE.exists():
        try: return json.loads(DATA_FILE.read_text(encoding="utf-8"))
        except: pass
    return {}

def save_projects(data):
    DATA_FILE.write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8")

def extract_docx(file_bytes):
    from docx import Document
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        tmp.write(file_bytes); tmp_path = tmp.name
    try:
        doc = Document(tmp_path)
        parts = [p.text for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                r = " | ".join(c.text.strip() for c in row.cells if c.text.strip())
                if r: parts.append(r)
        return "\n".join(parts)
    finally:
        try: os.unlink(tmp_path)
        except: pass

def extract_pdf(file_bytes):
    try:
        import pdfplumber
        with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as tmp:
            tmp.write(file_bytes); tmp_path = tmp.name
        try:
            with pdfplumber.open(tmp_path) as pdf:
                return "\n".join(p.extract_text() or "" for p in pdf.pages)
        finally:
            try: os.unlink(tmp_path)
            except: pass
    except ImportError:
        return "Cần cài pdfplumber: pip install pdfplumber"

TEMPLATE_TEXT = """HỢP ĐỒNG GIAO KHOÁN — MẪU CHUẨN ZALOPAY
Bên A: [Bên nhận khoán — Nhà thầu độc lập]
Bên B: CÔNG TY CỔ PHẦN ZION

Điều 1: Bên A là nhà thầu độc lập. Không phát sinh quan hệ lao động.
Điều 2: Thanh toán theo kết quả nghiệm thu sản phẩm/kết quả bàn giao. Không trả lương cố định hàng tháng.
Điều 3: Thời hạn hợp đồng tối đa 3 tháng.
Điều 4: Bên A tự chịu trách nhiệm BHXH tự nguyện và các bảo hiểm khác.
Điều 5: Bên B khấu trừ và nộp thuế TNCN cho cơ quan thuế trước khi thanh toán.
Điều 6: Quyền SHTT đối với kết quả công việc thuộc về Bên B.
Điều 7: Bên A cam kết bảo mật thông tin.
Điều 8: Phạt vi phạm 8% giá trị hợp đồng bị vi phạm.
Điều 9: Giải quyết tranh chấp tại Tòa án nhân dân TP.HCM.
"""

def review_contract(contract_text, filename, ref_mode="none", latest_text=None, ref_text=None):
    client = OpenAI(api_key=GN_API_KEY, base_url=GN_BASE_URL)
    system_prompt = load_system_prompt()

    if ref_mode == "template":
        ctx = "So sánh với Mẫu hợp đồng chuẩn Zalopay. Xác định các điều khoản khác mẫu và đánh giá rủi ro."
        ref_section = f"\n\n=== MẪU CHUẨN ZALOPAY (dùng làm Hợp đồng Mẫu) ===\n{TEMPLATE_TEXT}\n=== HẾT MẪU ==="
    elif ref_mode == "latest" and latest_text:
        ctx = "So sánh với bản Lastest đã được xác nhận gần nhất. Tập trung báo cáo những điểm thay đổi mới so với bản Lastest."
        ref_section = f"\n\n=== BẢN LASTEST THAM CHIẾU (Hợp đồng Mẫu) ===\n{latest_text}\n=== HẾT ==="
    elif ref_mode == "upload" and ref_text:
        ctx = "So sánh với file đối chiếu do người dùng upload. Xác định điểm khác biệt và đánh giá rủi ro."
        ref_section = f"\n\n=== FILE ĐỐI CHIẾU (Hợp đồng Mẫu) ===\n{ref_text}\n=== HẾT ==="
    else:
        ctx = "Rà soát độc lập — không so sánh với file gốc. Phân tích toàn bộ nội dung hợp đồng theo khung pháp lý."
        ref_section = ""

    content = f"File cần rà soát: {filename}\nChế độ: {ctx}\n\n=== NỘI DUNG HỢP ĐỒNG CẦN RÀ SOÁT ===\n{contract_text}\n=== HẾT ==={ref_section}"
    resp = client.chat.completions.create(
        model=GN_MODEL,
        max_tokens=4000,
        messages=[
            {"role": "system", "content": system_prompt},
            {"role": "user",   "content": content + "\n\nTrả kết quả theo đúng định dạng KẾT QUẢ 1 + KẾT QUẢ 2. Không có nội dung nào khác."}
        ]
    )
    return resp.choices[0].message.content

def review_edited(original_result, edited_text, mode):
    client = OpenAI(api_key=GN_API_KEY, base_url=GN_BASE_URL)
    system_prompt = load_system_prompt()
    if mode == "wording":
        prompt = f"Người dùng chỉnh sửa wording. Chỉ cập nhật phần diễn đạt, KHÔNG phân tích lại rủi ro:\n\nBản gốc:\n{original_result}\n\nBản chỉnh sửa:\n{edited_text}\n\nTrả về bản cập nhật wording, giữ nguyên cấu trúc bảng."
    else:
        prompt = f"Người dùng chỉnh sửa nội dung. Đọc lại và cập nhật đánh giá rủi ro:\n\nKết quả gốc:\n{original_result}\n\nNội dung sau khi sửa:\n{edited_text}\n\nTrả về kết quả rà soát đầy đủ với đánh giá rủi ro được cập nhật."
    resp = client.chat.completions.create(
        model=GN_MODEL,
        max_tokens=4000,
        messages=[
            {"role": "system", "content": system_prompt},
            {"role": "user",   "content": prompt}
        ]
    )
    return resp.choices[0].message.content

def load_html():
    """Đọc index.html từ cùng folder với app.py"""
    html_file = SCRIPT_DIR / "index.html"
    if html_file.exists():
        return html_file.read_text(encoding="utf-8")
    return "<h1>Không tìm thấy index.html — đặt file index.html cùng thư mục với app.py</h1>"

class Handler(BaseHTTPRequestHandler):
    def log_message(self, fmt, *args): print(f"[{self.address_string()}] {fmt % args}")

    def send_json(self, data, status=200):
        body = json.dumps(data, ensure_ascii=False).encode("utf-8")
        self.send_response(status)
        self.send_header("Content-Type", "application/json; charset=utf-8")
        self.send_header("Content-Length", str(len(body)))
        self.end_headers(); self.wfile.write(body)

    def do_GET(self):
        if self.path in ("/", "/index.html"):
            body = load_html().encode("utf-8")
            self.send_response(200)
            self.send_header("Content-Type", "text/html; charset=utf-8")
            self.send_header("Content-Length", str(len(body)))
            self.end_headers(); self.wfile.write(body)
        elif self.path == "/health":
            self.send_json({"ok": True})
        elif self.path == "/api/projects":
            self.send_json(load_projects())
        else:
            self.send_response(404); self.end_headers()

    def do_POST(self):
        ct = self.headers.get("Content-Type", "")
        cl = int(self.headers.get("Content-Length", 0))
        body = self.rfile.read(cl)

        if self.path == "/api/projects/create":
            data = json.loads(body)
            projects = load_projects()
            pid = f"proj_{int(time.time())}_{data['name'][:20].replace(' ','_')}"
            projects[pid] = {"id": pid, "name": data["name"], "created": data.get("created",""), "reviews": []}
            save_projects(projects)
            self.send_json({"id": pid, "project": projects[pid]})

        elif self.path == "/api/projects/delete":
            data = json.loads(body)
            projects = load_projects()
            projects.pop(data["id"], None)
            save_projects(projects)
            self.send_json({"ok": True})

        elif self.path == "/api/review":
            if "multipart/form-data" not in ct:
                self.send_json({"error": "Cần multipart/form-data"}, 400); return
            boundary = ct.split("boundary=")[-1].strip().encode()
            parts = body.split(b"--" + boundary)

            file_bytes = None; filename = "contract.docx"; project_id = ""
            ref_mode = "none"; latest_text = ""; ref_file_bytes = None

            for part in parts:
                if b"\r\n\r\n" not in part: continue
                hdr, pcontent = part.split(b"\r\n\r\n", 1)
                hdr_str = hdr.decode("utf-8", errors="replace")
                if pcontent.endswith(b"\r\n"): pcontent = pcontent[:-2]
                val = pcontent.decode("utf-8", "replace").strip()

                if 'name="project_id"' in hdr_str: project_id = val
                elif 'name="ref_mode"' in hdr_str: ref_mode = val
                elif 'name="latest_text"' in hdr_str: latest_text = val
                elif 'name="ref_file"' in hdr_str and 'filename="' in hdr_str:
                    ref_file_bytes = pcontent
                elif 'name="file"' in hdr_str and 'filename="' in hdr_str:
                    fn_s = hdr_str.find('filename="') + 10
                    filename = hdr_str[fn_s:hdr_str.find('"', fn_s)]
                    file_bytes = pcontent

            if not file_bytes:
                self.send_json({"error": "Không tìm thấy file"}, 400); return
            try:
                # Extract main file
                if filename.lower().endswith(".docx"):
                    text = extract_docx(file_bytes)
                elif filename.lower().endswith(".pdf"):
                    text = extract_pdf(file_bytes)
                else:
                    self.send_json({"error": "Chỉ hỗ trợ .docx và .pdf"}, 400); return

                if not text or len(text.strip()) < 30:
                    self.send_json({"error": "Không đọc được nội dung file"}); return

                # Extract ref file if uploaded
                ref_text = None
                if ref_file_bytes and ref_mode == "upload":
                    try:
                        if ref_file_bytes[:4] == b'PK\x03\x04':
                            ref_text = extract_docx(ref_file_bytes)
                        else:
                            ref_text = extract_pdf(ref_file_bytes)
                    except: pass

                # Determine latest_text usage
                lt = latest_text if ref_mode == "latest" else None

                print(f"Rà soát: {filename} | mode: {ref_mode}")
                result = review_contract(text, filename, ref_mode, lt, ref_text)

                # Save review
                projects = load_projects()
                if project_id in projects:
                    import base64 as _b64
                    rev = {
                        "id": f"rev_{int(time.time())}",
                        "filename": filename,
                        "timestamp": time.strftime("%d/%m/%Y %H:%M"),
                        "result": result,
                        "file_text": text,
                        "file_bytes_b64": _b64.b64encode(file_bytes).decode("utf-8"),
                        "ref_mode": ref_mode,
                        "confirmed": False
                    }
                    projects[project_id]["reviews"].append(rev)
                    save_projects(projects)
                    self.send_json({"result": result, "review_id": rev["id"]})
                else:
                    self.send_json({"result": result, "review_id": None})

            except Exception as e:
                import traceback; traceback.print_exc()
                self.send_json({"error": str(e)}, 500)

        elif self.path == "/api/review/edit":
            data = json.loads(body)
            try:
                result = review_edited(data["original"], data["edited"], data["mode"])
                projects = load_projects()
                pid, rid = data.get("project_id"), data.get("review_id")
                if pid and rid and pid in projects:
                    for r in projects[pid]["reviews"]:
                        if r["id"] == rid: r["result"] = result; break
                    save_projects(projects)
                self.send_json({"result": result})
            except Exception as e:
                self.send_json({"error": str(e)}, 500)

        elif self.path == "/api/review/confirm":
            data = json.loads(body)
            projects = load_projects()
            pid, rid = data.get("project_id"), data.get("review_id")
            if pid and rid and pid in projects:
                for r in projects[pid]["reviews"]:
                    if r["id"] == rid:
                        r["confirmed"] = True
                        r["confirmed_result"] = data.get("result", r["result"])
                        break
                save_projects(projects)
            self.send_json({"ok": True})

        elif self.path == "/api/owner-ai":
            data = json.loads(body)
            n         = data.get("n", "")
            dieu      = data.get("dieu", "")
            noi_dung  = data.get("noi_dung", "")
            phuong_an = data.get("phuong_an", "")
            risk_text = data.get("risk_text", "")
            y_kien    = data.get("y_kien", "")
            try:
                client = OpenAI(api_key=GN_API_KEY, base_url=GN_BASE_URL)
                resp = client.chat.completions.create(
                    model=GN_MODEL_JSON,
                    max_tokens=500,
                    extra_body={"enable_thinking": False, "chat_template_kwargs": {"enable_thinking": False}},
                    messages=[{
                        "role": "system",
                        "content": "You are a legal expert. Respond with ONLY a single JSON object on one line. No thinking, no explanation, no markdown. Just the JSON."
                    }, {
                        "role": "user",
                        "content": f"""Contract clause: {dieu}
Issue: {noi_dung}
AI proposal: {phuong_an}
Original risk: {risk_text}
Owner opinion: {y_kien}

Does the owner's opinion resolve the issue? Reply ONLY this JSON (no other text):
{{"risk":"ĐÃ XỬ LÝ or THẤP or TRUNG BÌNH or NGHIÊM TRỌNG","note":"one sentence in Vietnamese"}}"""
                    }]
                )
                import json as _json
                msg = resp.choices[0].message
                # Debug: print full response
                print(f"[owner-ai] model={GN_MODEL_JSON}")
                print(f"[owner-ai] content={repr(msg.content)}")
                print(f"[owner-ai] full_msg={msg}")
                
                # Read content — Qwen thinking mode puts answer in reasoning field
                raw = ""
                full = resp.model_dump()
                msg_dict = full.get("choices",[{}])[0].get("message",{})
                
                # Priority: content > reasoning_content > reasoning
                raw = (msg_dict.get("content") or "").strip()
                if not raw:
                    raw = (msg_dict.get("reasoning_content") or "").strip()
                if not raw:
                    raw = (msg_dict.get("reasoning") or "").strip()
                
                print(f"[owner-ai] raw_len={len(raw)}, source={'content' if msg_dict.get('content') else 'reasoning'}")
                
                raw = raw.replace("```json","").replace("```","").strip()
                if "{" in raw and "}" in raw:
                    raw = raw[raw.index("{"):raw.rindex("}")+1]
                print(f"[owner-ai] parsed_raw={repr(raw[:200])}")
                try:
                    result = _json.loads(raw)
                except Exception:
                    result = {"risk": data.get("risk_text","TRUNG BÌNH"), "note": raw[:150] if raw else "Không nhận được phản hồi từ AI"}
                self.send_json(result)
            except Exception as e:
                import traceback; traceback.print_exc()
                self.send_json({"error": str(e)}, 500)

        elif self.path == "/api/export":
            data = json.loads(body)
            export_type = data.get("type", "eval")
            project_id  = data.get("project_id", "")
            review_id   = data.get("review_id", "")
            owner_notes = data.get("owner_notes", {})

            projects = load_projects()
            proj = projects.get(project_id)
            rev  = None
            if proj:
                rev = next((r for r in proj["reviews"] if r["id"] == review_id), None)

            try:
                from docx import Document
                from docx.shared import Pt, RGBColor, Inches
                from docx.enum.text import WD_COLOR_INDEX, WD_ALIGN_PARAGRAPH
                from docx.oxml.ns import qn
                from docx.oxml import OxmlElement
                import io, re as _re

                def set_cell_bg(cell, hex_color):
                    tc = cell._tc
                    tcPr = tc.get_or_add_tcPr()
                    shd = OxmlElement("w:shd")
                    shd.set(qn("w:val"), "clear")
                    shd.set(qn("w:color"), "auto")
                    shd.set(qn("w:fill"), hex_color)
                    tcPr.append(shd)

                def add_cell(row, text, bold=False, bg=None, font_size=9, color=None):
                    cell = row.cells[len([c for c in row.cells if c.text or True])-1] if False else None
                    # use index
                    return text, bold, bg, font_size, color

                fname_map = {"eval": "Bang_danh_gia", "hd": "HD_da_update", "lastest": "Ban_Lastest"}
                fname = f"{fname_map.get(export_type,'export')}_{int(time.time())}.docx"

                # ── BẢNG ĐÁNH GIÁ ────────────────────────────────────────
                if export_type == "eval":
                    # ── BẢNG ĐÁNH GIÁ — giống màn hình ──────────────────
                    doc = Document()
                    section = doc.sections[0]
                    section.page_width  = Inches(13.0)
                    section.page_height = Inches(8.5)
                    section.left_margin = section.right_margin = Inches(0.5)
                    section.top_margin  = section.bottom_margin = Inches(0.5)

                    # Title
                    t = doc.add_paragraph()
                    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    tr = t.add_run("BẢNG ĐÁNH GIÁ HỢP ĐỒNG GIAO KHOÁN")
                    tr.bold = True; tr.font.size = Pt(14)
                    info = doc.add_paragraph(f"Project: {proj['name'] if proj else ''}  |  File: {rev['filename'] if rev else ''}  |  Ngày: {rev['timestamp'] if rev else ''}")
                    info.runs[0].font.size = Pt(9)
                    doc.add_paragraph()

                    result_text = rev.get("confirmed_result") or rev.get("result","") if rev else ""

                    # Parse markdown table rows
                    lines = result_text.splitlines()
                    table_lines = [l for l in lines if l.strip().startswith("|") and not all(c in "|-: " for c in l.strip())]

                    # Columns to show (matching screen): STT | Nhóm | Mức độ | Điều khoản | Nội dung | Phương án | Chiến lược | Ý kiến Owner | Phân quyền | PB tham vấn
                    # AI output columns: STT|Nhóm|Mức độ|Điều khoản|Nội dung|Đánh giá|Phương án|Chiến thuật|Phòng ban
                    # We skip col 5 (Đánh giá), keep: 0,1,2,3,4,6,7,8
                    KEEP_COLS   = [0,1,2,3,4,6,7,8]
                    HDR_LABELS  = ["#","Nhóm","Mức độ","Điều khoản","Nội dung cần lưu ý","Phương án xử lý","Chiến lược đàm phán","Phòng ban"]
                    COL_WIDTHS  = [400,900,900,1400,2200,2200,2000,900]  # in Twips

                    from docx.shared import Twips
                    ncols = len(HDR_LABELS)
                    tbl = doc.add_table(rows=1, cols=ncols)
                    tbl.style = "Table Grid"
                    tbl.autofit = False

                    # Header
                    hdr = tbl.rows[0]
                    for i,(lbl,w) in enumerate(zip(HDR_LABELS,COL_WIDTHS)):
                        c = hdr.cells[i]; c.width = Twips(w)
                        set_cell_bg(c,"0B1E3D")
                        r = c.paragraphs[0].add_run(lbl)
                        r.bold=True; r.font.size=Pt(8)
                        r.font.color.rgb=RGBColor(255,255,255)

                    RISK_BG = {"NGHIÊM":"FFF0F0","TRỌNG":"FFF0F0","TRUNG BÌNH":"FFFBF0","THẤP":"F0FFF4"}

                    row_num = 0
                    for line in table_lines[1:]:
                        parts = [c.strip() for c in line.split("|")]
                        if len(parts) < 2: continue
                        # Remove empty edge cells from | split
                        if parts and parts[0]=="": parts=parts[1:]
                        if parts and parts[-1]=="": parts=parts[:-1]
                        if not any(p for p in parts): continue

                        # Determine row background from risk level
                        risk_str = " ".join(parts).upper()
                        bg = "FFFFFF"
                        if "NGHIÊM" in risk_str or "🔴" in risk_str: bg="FFF0F0"
                        elif "TRUNG" in risk_str or "🟡" in risk_str: bg="FFFBF0"
                        elif "THẤP" in risk_str or "🟢" in risk_str: bg="F0FFF4"

                        row_num += 1
                        new_row = tbl.add_row()
                        for col_idx, keep in enumerate(KEEP_COLS):
                            val = parts[keep].strip() if keep < len(parts) else ""
                            # Strip leading (1)(2)... from phương án col
                            if col_idx == 5:
                                import re as _re2
                                val = _re2.sub(r'^(\(\d+\)\s*)+','',val).strip()
                            c2 = new_row.cells[col_idx]
                            c2.width = Twips(COL_WIDTHS[col_idx])
                            set_cell_bg(c2, bg)
                            c2.paragraphs[0].add_run(val).font.size = Pt(8)

                    # Owner notes section
                    if owner_notes:
                        doc.add_paragraph()
                        h = doc.add_paragraph()
                        h.add_run("Ý KIẾN & ĐÁNH GIÁ LẠI CỦA OWNER").bold = True
                        for k,v in owner_notes.items():
                            p2 = doc.add_paragraph(style="List Bullet")
                            p2.add_run(f"Dòng {k}: ").bold = True
                            p2.add_run(v.get("note","")).font.size = Pt(9)
                            if v.get("ai_result"):
                                p3 = doc.add_paragraph()
                                p3.add_run(f"  → AI đánh giá lại: {v['ai_result']}").italic = True

                elif export_type == "hd":
                    import base64 as _b64hd, tempfile as _tmp
                    # Use original uploaded docx file bytes
                    file_b64 = rev.get("file_bytes_b64","") if rev else ""
                    file_text = (rev.get("file_text") or "") if rev else ""
                    result_text = (rev.get("confirmed_result") or rev.get("result","")) if rev else ""

                    # Parse issues from result
                    issue_lines=[l for l in result_text.splitlines() if l.strip().startswith("|") and "---" not in l]
                    import re as _re3
                    issues=[]
                    for line in (issue_lines[1:] if len(issue_lines)>1 else []):
                        cols=[c.strip() for c in line.split("|")][1:-1]
                        if len(cols)>=6:
                            pa=_re3.sub(r"[(][0-9][)]\s*","",cols[5] if len(cols)>5 else "").strip()
                            issues.append({"risk":cols[1] if len(cols)>1 else "","dieu":cols[2] if len(cols)>2 else "","noidung":cols[3] if len(cols)>3 else "","phuongan":pa})

                    if file_b64:
                        # Use original docx and modify it
                        raw_bytes = _b64hd.b64decode(file_b64)
                        with _tmp.NamedTemporaryFile(suffix=".docx", delete=False) as tf:
                            tf.write(raw_bytes); tf_path = tf.name
                        from docx import Document as DocxDoc
                        from docx.oxml.ns import qn
                        from docx.oxml import OxmlElement
                        from docx.shared import RGBColor as RC2, Pt as Pt2
                        from docx.enum.text import WD_COLOR_INDEX as WCI2
                        import copy

                        doc = DocxDoc(tf_path)
                        try: os.unlink(tf_path)
                        except: pass

                        def para_matches(para_text, issue):
                            pl = para_text.lower()
                            kw = [w for w in issue["noidung"].lower().split() if len(w)>4][:4]
                            dl = issue["dieu"].lower()
                            return (dl and dl[:15] in pl) or (kw and sum(1 for w in kw if w in pl)>=2)

                        def strike_run(run):
                            rpr = run._r.get_or_add_rPr()
                            strike = OxmlElement("w:strike")
                            strike.set(qn("w:val"), "true")
                            rpr.append(strike)

                        def add_annotation(para, issue):
                            is_sev = "NGHIEM" in issue["risk"].upper() or "NGHIÊM" in issue["risk"].upper()
                            is_med = "TRUNG" in issue["risk"].upper()
                            ri = "🔴" if is_sev else "🟡" if is_med else "🟢"
                            new_para = OxmlElement("w:p")
                            new_run = OxmlElement("w:r")
                            rpr = OxmlElement("w:rPr")
                            color_el = OxmlElement("w:color")
                            color_el.set(qn("w:val"), "A32D2D" if is_sev else "854F0B")
                            sz = OxmlElement("w:sz"); sz.set(qn("w:val"), "18")
                            italic_el = OxmlElement("w:i")
                            rpr.append(color_el); rpr.append(sz); rpr.append(italic_el)
                            new_run.append(rpr)
                            t = OxmlElement("w:t")
                            t.text = f"  >> {ri} [{issue['dieu']}] De xuat: {issue['phuongan']}"
                            t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
                            new_run.append(t); new_para.append(new_run)
                            para._p.addnext(new_para)

                        # Process each paragraph
                        for para in doc.paragraphs:
                            if not para.text.strip(): continue
                            for issue in issues:
                                if para_matches(para.text, issue):
                                    # Highlight yellow + strikethrough original
                                    for run in para.runs:
                                        run.font.highlight_color = WCI2.YELLOW
                                    # Add suggestion paragraph after
                                    add_annotation(para, issue)
                                    # Add replacement run
                                    if issue["phuongan"]:
                                        new_run2 = para.add_run(" → " + issue["phuongan"])
                                        new_run2.bold = True
                                        new_run2.font.size = Pt2(10)
                                        new_run2.font.color.rgb = RC2(0x08,0x50,0x41)
                                    break

                        # Process tables too
                        for table in doc.tables:
                            for row in table.rows:
                                for cell in row.cells:
                                    for para in cell.paragraphs:
                                        if not para.text.strip(): continue
                                        for issue in issues:
                                            if para_matches(para.text, issue):
                                                for run in para.runs:
                                                    run.font.highlight_color = WCI2.YELLOW
                                                break

                        # Owner notes at end
                        if owner_notes:
                            doc.add_paragraph()
                            h2p = doc.add_paragraph()
                            h2p.add_run("GHI CHU OWNER & AI DANH GIA LAI").bold = True
                            for k, v in owner_notes.items():
                                p2 = doc.add_paragraph()
                                r1 = p2.add_run("Dong "+str(k)+": "); r1.bold=True; r1.font.size=Pt2(9)
                                p2.add_run(v.get("note","")).font.size=Pt2(9)
                                if v.get("ai_result"):
                                    p3=doc.add_paragraph()
                                    r3=p3.add_run("  -> AI: "+v["ai_result"])
                                    r3.font.size=Pt2(9); r3.italic=True
                                    r3.font.color.rgb=RC2(0x08,0x50,0x41)
                    else:
                        # Fallback: create from text
                        doc = Document()
                        section = doc.sections[0]
                        section.left_margin=section.right_margin=Inches(1.0)
                        section.top_margin=section.bottom_margin=Inches(1.0)
                        doc.add_paragraph().add_run("HOP DONG DA CAP NHAT").bold=True
                        for line in file_text.splitlines():
                            if not line.strip(): doc.add_paragraph()
                            else: doc.add_paragraph().add_run(line).font.size=Pt(11)

                elif export_type == "lastest":
                    # ── BẢN LASTEST ──────────────────────────────────────
                    doc = Document()
                    section = doc.sections[0]
                    section.left_margin = section.right_margin = Inches(1.0)
                    section.top_margin  = section.bottom_margin = Inches(1.0)

                    t = doc.add_paragraph()
                    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    t.add_run("BẢN LASTEST — HỢP ĐỒNG GIAO KHOÁN").bold = True

                    conf = [r for r in (proj["reviews"] if proj else []) if r.get("confirmed")]
                    last = conf[-1] if conf else rev
                    doc.add_paragraph(f"Project: {proj['name'] if proj else ''}")
                    doc.add_paragraph(f"Xác nhận ngày: {last['timestamp'] if last else ''}")
                    doc.add_paragraph(f"File gốc: {last['filename'] if last else ''}")
                    doc.add_paragraph()
                    ft = last.get("file_text","") if last else ""
                    for para_text in ft.splitlines():
                        if not para_text.strip(): doc.add_paragraph()
                        else: doc.add_paragraph().add_run(para_text).font.size = Pt(11)
                # Save
                buf = io.BytesIO()
                doc.save(buf)
                buf.seek(0)
                docx_bytes = buf.read()

                self.send_response(200)
                self.send_header("Content-Type", "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
                self.send_header("Content-Disposition", f'attachment; filename="{fname}"')
                self.send_header("Content-Length", str(len(docx_bytes)))
                self.end_headers()
                self.wfile.write(docx_bytes)

            except Exception as e:
                import traceback; traceback.print_exc()
                self.send_json({"error": str(e)}, 500)

        else:
            self.send_response(404); self.end_headers()

def main():
    print(f"\nZalopay Legal Contract — Agent\nhttp://localhost:{PORT}\n")
    server = HTTPServer(("0.0.0.0", PORT), Handler)
    try: server.serve_forever()
    except KeyboardInterrupt: server.shutdown()

if __name__ == "__main__": main()
